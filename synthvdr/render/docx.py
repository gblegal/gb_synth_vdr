"""Optional DOCX render. Never imported at core-build time.

`render_tree_docx` deliberately does NOT delete-and-rebuild its output
directory the way `synthvdr.twin`, `synthvdr.subset` and
`synthvdr.index_build` do. All three of those shipped a defect where they
destroyed data they did not own before the shared ownership guard in
`synthvdr.ownership` closed it — the cheapest fix to a destructive site is
one that never exists. This writer only ever creates directories
(`mkdir(parents=True, exist_ok=True)`) and (re)writes the `.docx` files it
is responsible for; it never removes a file or directory, including a
render whose source document was since renamed or deleted. That kind of
staleness is a corpus-consistency problem, not a rendering problem, and is
caught instead by `synthvdr.qa.renders.gate_16_render_parity`, checking in
both directions: sources missing a render, and renders missing a source.
"""

from __future__ import annotations

import csv
import hashlib
import re
from pathlib import Path
from typing import List, Optional

from ..schema import FindingSet


class RenderUnavailable(Exception):
    """A render toolchain is not installed."""


# CommonMark requires whitespace between the '#' run and the heading text —
# `#MeToo` or `#1 supplier` are NOT headings in any markdown dialect, they
# are paragraphs that happen to start with a hash. A bare `stripped.startswith
# ("#")` check (this module's original implementation) got that wrong, and
# `lstrip("# ")` compounded it by eating any further leading '#'/' ' pairs —
# a heading whose own title legitimately starts with '#' (e.g. "# #1
# priority") lost that leading character too. Bounded at 6 hashes, per
# CommonMark ATX headings — a 7th leading '#' means the line has no opening
# sequence at all, and stays a plain paragraph, hashes and all.
#
# The separator is `[ \t]+` — ASCII space and tab, EXACTLY the two
# characters CommonMark's ATX heading spec names — not `\s+`. `\s` is
# Unicode-wide: it also matches NBSP (U+00A0, which copy-pasted prose
# carries constantly), vertical tab, form feed, EN SPACE (U+2002) and
# IDEOGRAPHIC SPACE (U+3000, live in a corpus that already contemplates CJK
# documents elsewhere in this harness). Under `\s+`, "#\xa0Title" silently
# became a heading — the exact corruption class this whole heading fix
# exists to close, reopened one character class wider. This has now gone
# wrong twice in this function alone (`startswith("#")` too loose one way,
# `\s` too loose another): name the exact separator set the spec names,
# never the one that merely reads naturally.
_ATX_HEADING = re.compile(r"^(#{1,6})[ \t]+(.*)$")

# A fenced code block (``` or ~~~, CommonMark's two fence characters) is
# tracked across lines so nothing inside it is ever mistaken for a heading
# — "# a shell comment" is the single most common line in any shell or
# Python snippet, and it is not a heading just because it starts a fresh
# line inside a fence. Every line inside a fence, including the fence
# markers themselves, is rendered as a verbatim paragraph: this project's
# rule is that content survives the render, and dropping the fence marker
# lines to make the output prettier is exactly the trade this task exists
# to refuse. An unclosed fence (no matching closing marker before EOF)
# falls out of the same state machine for free: every remaining line stays
# in fenced mode through to the end of the document.
_FENCE = re.compile(r"^(`{3,}|~{3,})")


def rotation_for(slot_id: str, page: int) -> float:
    """Deterministic skew in +/- 0.4-1.1 degrees, derived from sha256 of the
    slot id and page number. No RNG, no bare `hash()`, no clock: the same
    (slot_id, page) pair must produce the same angle on every run, in every
    process, regardless of `PYTHONHASHSEED` — sha256 does not depend on it,
    which is precisely why it is used here instead of Python's built-in
    `hash()`, whose string hashing IS seed-salted.
    """
    digest = hashlib.sha256(f"{slot_id}:{page}".encode("utf-8")).digest()
    magnitude = 0.4 + (digest[0] / 255) * 0.7
    return magnitude if digest[1] % 2 == 0 else -magnitude


def scanned_slots(findings: FindingSet, count: int, *, suffix: Optional[str] = None) -> List[str]:
    """Pick `count` scanned slots, drawn only from evidence documents — the
    OCR challenge belongs where the substance is, so OCR failure costs the
    tool a finding. Never backfilled with non-evidence documents: if
    `count` exceeds the number of distinct evidence paths, fewer than
    `count` slots are returned rather than diluting the bias.

    Deterministic and reproducible across processes and across reorderings
    of `findings`: `all_evidence_paths()` returns a set, so any incoming
    order collapses through it, and the two sorts below (first
    lexicographic, then by sha256 digest) fix one specific, seed-independent
    order out of that set every time.

    `suffix`, when given, narrows the pool BEFORE selection rather than
    filtering the result afterwards — filtering after would silently return
    fewer than `count` for a reason that has nothing to do with the
    "never backfill" rule above, and the two shortfalls must not be
    confusable. `write_scanned_csv` passes ".md" because the renderers only
    produce a page per markdown source; a CSV register named as evidence is
    real evidence with no page to scan.
    """
    evidence = sorted(findings.all_evidence_paths())
    if suffix is not None:
        evidence = [path for path in evidence if path.endswith(suffix)]
    ordered = sorted(evidence, key=lambda p: hashlib.sha256(p.encode("utf-8")).hexdigest())
    return ordered[:count]


# The share of a room's markdown evidence documents that ships as a scan.
# A function below rather than a number quoted in prose in /vdr-package, for
# the same reason synthvdr.schema.severity_targets is a function: the moment
# the rule lives in a skill's prose, the skill and the code can disagree
# about it and nothing notices.
SCANNED_SHARE = 0.25


def default_scanned_count(findings: FindingSet) -> int:
    """How many evidence documents a room of this size should ship as scans.

    A quarter of the markdown evidence, rounded, and never zero while there
    is anything to scan — a room with no scanned page at all does not test
    OCR, which is the one thing the PDF render exists to add over the
    markdown a tool could otherwise read directly. Returns 0 only when there
    is genuinely no markdown evidence to draw from.
    """
    pool = [path for path in findings.all_evidence_paths() if path.endswith(".md")]
    if not pool:
        return 0
    return max(1, round(len(pool) * SCANNED_SHARE))


def write_scanned_csv(findings: FindingSet, count: int, path: Path) -> List[str]:
    """Write the `slot,page` manifest `pdf.mjs` reads, and return the slots.

    THIS IS THE STEP THAT WAS MISSING. `scanned_slots` and `pdf.mjs`'s
    `loadScannedSlots` were both written, both tested, and never connected:
    nothing in this package or in any skill ever produced the file, so
    `loadScannedSlots` returned an empty map on every real run and no room
    ever shipped a scanned page — while README and TECHNICAL-NOTES §5 both
    described the feature as if it did.

    `slot` is the source path relative to BLIND_TREE with `.md` stripped and
    forward slashes, exactly as `pdf.mjs` reconstructs it from its own walk;
    a mismatch here is invisible (an unmatched slot is simply never scanned),
    which is why `test_scanned_csv_slots_match_pdf_mjs_slot_ids` derives the
    expectation from that file's own expression rather than restating it.

    Every row is page 1. `pdf.mjs` today honours only page 1 — it reads and
    stores every row and then asks `scannedPages.has(1)` — so writing a
    higher page number would produce a manifest line that silently does
    nothing. When that is fixed, this is where multi-page rows belong.
    """
    slots = scanned_slots(findings, count, suffix=".md")
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["slot", "page"])
        for rel in slots:
            writer.writerow([rel[: -len(".md")], 1])
    return slots


def render_tree_docx(src: Path, out: Path) -> int:
    """Render every markdown file under `src` to a `.docx` twin under `out`,
    mirroring `src`'s relative layout. Returns the number of files written.

    Non-destructive by design (see module docstring): creates directories
    as needed and writes/overwrites the `.docx` files it produces, but
    never deletes anything, including a stale render left behind by a
    since-renamed or since-deleted source.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RenderUnavailable(
            "python-docx is not installed. Install it with: pip install 'synthvdr[docx]'"
        ) from exc

    written = 0
    for path in sorted(src.rglob("*.md")):
        rel = path.relative_to(src).with_suffix(".docx")
        target = out / rel
        target.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        in_fence = False
        fence_char = None
        for line in path.read_text(encoding="utf-8").splitlines():
            stripped = line.rstrip()

            if in_fence:
                # Verbatim, no exceptions — including a line that would
                # otherwise look like a heading. The opening fence marker
                # went through this same append on the way in.
                document.add_paragraph(stripped)
                closing = _FENCE.match(stripped)
                if closing and closing.group(1)[0] == fence_char:
                    in_fence = False
                    fence_char = None
                continue

            fence = _FENCE.match(stripped)
            if fence:
                in_fence = True
                fence_char = fence.group(1)[0]
                document.add_paragraph(stripped)
                continue

            heading = _ATX_HEADING.match(stripped)
            if heading:
                # group(1) is exactly the matched hashes (1-6 of them, by
                # construction of the regex); group(2) is exactly what
                # follows the required whitespace — never a blanket
                # lstrip, so a title starting with '#' is preserved.
                level = min(len(heading.group(1)), 4)
                document.add_heading(heading.group(2), level=level)
            elif stripped:
                document.add_paragraph(stripped)
        document.save(str(target))
        written += 1
    return written
